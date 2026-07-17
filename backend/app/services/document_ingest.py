"""Ingest an uploaded file into the platform's knowledge base. Text is extracted from PDF,
Word, or plain-text files, a documents row is recorded, and the body is chunked and embedded
into document_embeddings — the same store the seeded corpus uses, so an uploaded file becomes
retrievable through the existing hybrid search and the copilot without any special-casing. The
original bytes are also saved to disk so the file itself can be downloaded again later.
"""

import asyncio
import os
import tempfile
from datetime import date
from io import BytesIO
from pathlib import Path

from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models import Document, DocumentEmbedding
from app.services.chunking import chunk_text
from app.services.embeddings import EmbeddingClient

MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
_TEXT_EXTENSIONS = (".txt", ".md", ".markdown", ".csv", ".log")
_ALL_EXTENSIONS = (".pdf", ".docx", *_TEXT_EXTENSIONS)


class UnsupportedDocument(ValueError):
    """Raised when a file's type is not one we can parse."""


class EmptyDocument(ValueError):
    """Raised when a supported file yields no extractable text (e.g. a scanned PDF)."""


def _extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in _ALL_EXTENSIONS:
        raise UnsupportedDocument("Unsupported file type. Upload a PDF, DOCX, or text file.")
    return ext


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(ext: str, data: bytes) -> str:
    if ext == ".pdf":
        try:
            return _extract_pdf(data)
        except Exception as exc:  # pypdf raises PdfReadError/PdfStreamError on malformed input
            raise UnsupportedDocument("The file's contents could not be read as a PDF.") from exc
    if ext == ".docx":
        try:
            return _extract_docx(data)
        except Exception as exc:  # python-docx raises on malformed/non-DOCX input
            raise UnsupportedDocument(
                "The file's contents could not be read as a Word document."
            ) from exc
    return data.decode("utf-8", errors="ignore")


def _upload_root() -> Path:
    """Root directory original upload files are saved under. Under TESTING this is the
    container's own ephemeral filesystem, not the real uploads volume — `docker compose run
    --rm` destroys it along with the container, so test runs never pollute real storage and
    need no manual cleanup, mirroring how embeddings.py/llm.py branch on TESTING."""
    if os.environ.get("TESTING"):
        return Path(tempfile.gettempdir()) / "construction-ai-test-uploads"
    return Path(get_settings().upload_dir)


def _save_upload(document_id: int, ext: str, data: bytes) -> str:
    root = _upload_root()
    root.mkdir(parents=True, exist_ok=True)
    path = root / f"{document_id}{ext}"
    path.write_bytes(data)
    return str(path)


async def ingest_upload(
    db: AsyncSession,
    embedder: EmbeddingClient,
    *,
    project_id: int,
    doc_type: str,
    title: str,
    filename: str,
    data: bytes,
) -> tuple[Document, int, int]:
    """Extract, persist, chunk, embed, and save the original file. Returns
    (document, chunk_count, character_count).

    Parsing runs in a worker thread because pypdf/python-docx are synchronous and a large
    file would otherwise block the event loop. The original file is written to disk as the
    LAST step, after the database row and its embeddings are already staged: if that write
    fails, the exception propagates before the caller's db.commit() ever runs, so the whole
    request rolls back cleanly with nothing orphaned in the database or on disk.
    """
    ext = _extension(filename)
    text = (await asyncio.to_thread(extract_text, ext, data)).strip()
    if not text:
        raise EmptyDocument("No readable text could be extracted from the file.")

    document = Document(
        project_id=project_id,
        doc_type=doc_type,
        title=title,
        doc_date=date.today(),
        content_summary=" ".join(text.split())[:500],
        original_filename=filename[:255],
    )
    db.add(document)
    await db.flush()  # assign document.id before referencing it from the chunks and file path

    chunks = chunk_text(text)
    vectors = await embedder.embed_documents(chunks)
    for index, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True)):
        db.add(
            DocumentEmbedding(
                source_type="document",
                source_id=document.id,
                project_id=project_id,
                chunk_index=index,
                content=chunk,
                token_count=max(len(chunk) // 4, 1),
                embedding=vector,
            )
        )

    document.storage_path = await asyncio.to_thread(_save_upload, document.id, ext, data)
    return document, len(chunks), len(text)
